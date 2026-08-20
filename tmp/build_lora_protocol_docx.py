from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("智能通风系统_LoRa报文协议_最终版.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
RED = "9B1C1C"
GOLD = "7A5A00"
WHITE = "FFFFFF"
FONT = "Microsoft YaHei"
MONO = "Consolas"
CONTENT_DXA = 9360


def set_run_font(run, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text, size=11, color=None, bold=False, italic=False, font=FONT):
    run = paragraph.add_run(text)
    set_run_font(run, name=font, size=size, color=color, bold=bold, italic=italic)
    return run


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=after, line=1.25)
    add_text(p, text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    add_text(p, text, size={1: 16, 2: 13, 3: 12}[level],
             color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level], bold=True)
    return p


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade(cell, "F7F9FB")
    set_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.15)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        add_text(p, line, size=9.5, color=NAVY, font=MONO)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, label, text, color=BLUE, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=0, line=1.2)
    add_text(p, f"{label}  ", size=11, color=color, bold=True)
    add_text(p, text, size=11, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.1)
        add_text(p, str(text), size=10, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.15)
            add_text(p, str(value), size=10)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)


def set_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=0, line=1.0)
    add_text(p, "智能通风系统  |  LoRa 报文协议（最终版）", size=9, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=0, line=1.0)
    add_text(p, "内部工程文档  |  V1.0  |  第 ", size=9, color=MUTED)
    add_page_field(p)
    add_text(p, " 页", size=9, color=MUTED)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    add_text(p, "智能通风系统", size=12, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_text(p, "LoRa 报文协议（最终版）", size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    add_text(p, "控制室网关、主机、从机统一执行标准", size=13, color=MUTED)

    add_callout(
        doc,
        "最终决定",
        "36 个温度点作为一套温度快照，一次 LoRa 报文完整发送；不再按 6 个 Port 拆分发送。",
        color=DARK_BLUE,
        fill=LIGHT_BLUE,
    )

    add_table(
        doc,
        ["项目", "固定要求"],
        [
            ("协议版本", "V1.0"),
            ("适用设备", "1 台控制室网关 G、4 台主机 M1-M4、4 台从机 S1-S4"),
            ("适用链路", "控制室 <-> 主机 <-> 同组从机"),
            ("温度传输", "36 点温度 + 36 位有效标志，一次 TEMP_36 报文发送"),
            ("LoRa 包长", "所有模块统一设为 128 B"),
            ("生效状态", "本文件为全员唯一报文与指令标准"),
        ],
        [2700, 6660],
    )

    add_heading(doc, "先看这一页：全员必须遵守", 1)
    add_body(doc, "1. 控制室是唯一调度者。主机和从机平时只监听，被点名后才回复。")
    add_body(doc, "2. 控制室不直接找从机；主机只找与自己拨码相同的从机。")
    add_body(doc, "3. 温度零值必须标记为无效。上位机显示“--”，不能当成 0℃参与控制。")
    add_body(doc, "4. 命令收到后先回 ACK，实际写入 Flash 和 TD710 成功后再回 RESULT。")
    add_body(doc, "5. 手动停止优先级最高，只有收到 SET_AUTO 后才允许恢复自动控制。")

    doc.add_page_break()


def add_main_content(doc):
    add_heading(doc, "1. 系统通信关系", 1)
    add_code_block(doc, [
        "控制室网关 G",
        "   |  点名、下发控制",
        "   v",
        "主机 M1 / M2 / M3 / M4",
        "   |  仅在本机温度缓存过期时请求",
        "   v",
        "同组从机 S1 / S2 / S3 / S4",
    ])
    add_body(doc, "每一组由一台主机和一台从机组成：M1 对应 S1，M2 对应 S2，以此类推。主机与从机拨码相同，才属于同一业务组。")
    add_callout(doc, "不允许", "从机周期性主动无线发送温度；非目标主机或从机抢占无线信道；控制室直接向从机发控制命令。", color=RED, fill="FDEEEE")

    add_heading(doc, "2. 角色、组号与基本单位", 1)
    add_table(
        doc,
        ["项目", "定义"],
        [
            ("控制室角色", "G，角色码 01，组号 00"),
            ("主机角色", "M，角色码 02，组号 01-04"),
            ("从机角色", "S，角色码 03，组号 01-04"),
            ("温度单位", "int16，单位 0.1℃。例如 258 表示 25.8℃"),
            ("频率单位", "uint16，单位 0.01 Hz。例如 5000 表示 50.00 Hz"),
            ("温度排序", "温度 1-6 为 Port 1；温度 7-12 为 Port 2；依次到 Port 6"),
        ],
        [2700, 6660],
    )

    add_heading(doc, "3. 所有报文的共同格式", 1)
    add_body(doc, "所有无线应用数据都采用二进制帧，不附加回车、换行或文本分隔符。用一句话概括，就是：")
    add_code_block(doc, ["【帧头 | 版本 | 报文类型 | 谁发的 | 发给谁 | 序号 | 数据长度 | 数据 | CRC16】"])
    add_table(
        doc,
        ["字节位置", "字段", "长度", "固定要求"],
        [
            ("0-1", "帧头", "2", "AA 55"),
            ("2", "版本", "1", "01"),
            ("3", "报文类型", "1", "见第 4 节"),
            ("4", "发送者角色", "1", "01=G，02=M，03=S"),
            ("5", "发送者组号", "1", "G 为 00；设备为 01-04"),
            ("6", "接收者角色", "1", "01=G，02=M，03=S"),
            ("7", "接收者组号", "1", "G 为 00；设备为 01-04"),
            ("8-9", "序号 seq", "2", "低字节在前；用于匹配、去重与重发"),
            ("10", "数据长度 len", "1", "数据区实际长度"),
            ("11 起", "数据 payload", "0-96", "按报文类型解释"),
            ("最后 2 字节", "CRC16", "2", "CRC-16/MODBUS，低字节在前"),
        ],
        [1200, 1900, 950, 5310],
    )
    add_callout(doc, "长度规则", "完整应用报文 = 13 + len 字节。TEMP_36 的 len=81，完整报文为 94 字节，小于统一配置的 128 B LoRa 包长。", color=DARK_BLUE, fill=LIGHT_BLUE)

    add_heading(doc, "4. 报文类型速查表", 1)
    add_table(
        doc,
        ["类型", "名称", "谁发给谁", "一句话说明"],
        [
            ("01", "READ_TEMP", "G -> M；M -> S", "请求一整套 36 点温度"),
            ("02", "TEMP_36", "S -> M；M -> G", "一次返回全部 36 点温度"),
            ("10", "CMD", "G -> M", "设置频率、温度、手动启动/停止、切自动"),
            ("11", "ACK", "M -> G", "主机已收到控制命令"),
            ("12", "RESULT", "M -> G", "命令真实执行结果"),
            ("7E", "ERROR", "双向", "处理失败、组号错误、超时等"),
        ],
        [900, 1700, 2600, 4160],
    )

    add_heading(doc, "5. 温度报文：一次发送全部 36 点", 1)
    add_heading(doc, "5.1 READ_TEMP：请求温度", 2)
    add_table(
        doc,
        ["数据字段", "长度", "说明"],
        [
            ("请求编号 request_id", "2", "控制室生成；主机转发给从机时原样带上"),
            ("读取模式", "1", "00=允许主机使用新鲜缓存；01=强制从机重新采样"),
        ],
        [3300, 1000, 5060],
    )
    add_code_block(doc, [
        "G -> M1：READ_TEMP，request_id=100，读取模式=00",
        "意思：控制室请 M1 发一套 36 点温度；缓存新鲜就直接返回。",
    ])

    add_heading(doc, "5.2 TEMP_36：完整 36 点温度", 2)
    add_table(
        doc,
        ["数据字段", "长度", "说明"],
        [
            ("请求编号 request_id", "2", "这套温度属于哪一次 READ_TEMP 请求"),
            ("数据来源", "1", "01=从机本次新采样；02=主机缓存"),
            ("快照状态", "1", "00=完整；01=不完整；02=采样失败"),
            ("有效位图 valid_mask", "5", "36 个温度一一对应；第 5 字节高 4 位固定为 0"),
            ("温度 1 到温度 36", "72", "每个温度占 2 字节，int16，单位 0.1℃"),
        ],
        [3300, 1000, 5060],
    )
    add_callout(doc, "零值与无效值", "温度数据为 0 时，对应有效位必须为 0。上位机显示“--”；该点不得参与“全 36 点低温停机”判断。", color=RED, fill="FDEEEE")
    add_body(doc, "有效位图从低位开始对应温度 1、温度 2、温度 3……温度 36。有效位为 1 表示该温度可信；为 0 表示无效。")
    add_table(
        doc,
        ["温度编号", "来源"],
        [
            ("温度 1-6", "Port 1 的 6 个温度"),
            ("温度 7-12", "Port 2 的 6 个温度"),
            ("温度 13-18", "Port 3 的 6 个温度"),
            ("温度 19-24", "Port 4 的 6 个温度"),
            ("温度 25-30", "Port 5 的 6 个温度"),
            ("温度 31-36", "Port 6 的 6 个温度"),
        ],
        [2200, 7160],
    )

    add_heading(doc, "6. 控制命令与返回结果", 1)
    add_heading(doc, "6.1 CMD：控制室发给指定主机", 2)
    add_table(
        doc,
        ["数据字段", "长度", "说明"],
        [
            ("命令号", "1", "见下表"),
            ("参数 value", "2", "频率或温度命令使用；其余命令为 0"),
        ],
        [3300, 1000, 5060],
    )
    add_table(
        doc,
        ["命令号", "命令名称", "参数含义", "主机必须执行的动作"],
        [
            ("01", "SET_FREQ", "0.01 Hz", "保存频率；写入 TD710；成功后才返回 RESULT 成功"),
            ("02", "SET_TARGET_TEMP", "0.1℃", "保存目标温度"),
            ("03", "MANUAL_RUN", "0", "切到手动运行并请求风机运行"),
            ("04", "MANUAL_STOP", "0", "切到手动停止并请求风机停止；保持停机"),
            ("05", "SET_AUTO", "0", "退出手动状态，恢复自动温控"),
            ("06", "QUERY_STATUS", "0", "返回主机模式、风机状态、频率和目标温度"),
        ],
        [900, 2050, 1500, 4910],
    )
    add_callout(doc, "手动停止规则", "MANUAL_STOP 优先级最高。即使任一温度超温，也不得自动启动；只有 SET_AUTO 能解除该锁定。", color=RED, fill="FDEEEE")

    add_heading(doc, "6.2 ACK：主机已收到，不等于执行成功", 2)
    add_table(
        doc,
        ["数据字段", "长度", "说明"],
        [
            ("原命令序号", "2", "对应 CMD 的 seq"),
            ("接收状态", "1", "00=接受；01=拒绝"),
            ("拒绝原因", "1", "00=无；其余见 ERROR/RESULT 错误码"),
        ],
        [3300, 1000, 5060],
    )
    add_heading(doc, "6.3 RESULT：实际执行结果", 2)
    add_table(
        doc,
        ["数据字段", "长度", "说明"],
        [
            ("原命令序号", "2", "对应 CMD 的 seq"),
            ("执行结果", "1", "00=成功；其他值见下表"),
            ("当前模式", "1", "00=AUTO；01=MANUAL_RUN；02=MANUAL_STOP；03=FAULT"),
            ("风机状态", "1", "00=停止；01=运行；02=未知"),
            ("当前频率", "2", "单位 0.01 Hz"),
            ("当前目标温度", "2", "单位 0.1℃"),
        ],
        [3300, 1000, 5060],
    )
    add_table(
        doc,
        ["结果码", "含义", "上位机显示"],
        [
            ("00", "成功", "设置成功 / 已启动 / 已停止"),
            ("01", "参数范围错误", "参数错误"),
            ("02", "当前状态不允许", "命令被拒绝"),
            ("03", "Flash 保存失败", "参数未保存"),
            ("04", "TD710 无响应超时", "变频器无响应"),
            ("05", "TD710 返回错误", "变频器执行失败"),
            ("06", "重复命令", "返回上一次结果，不重复执行"),
            ("07", "主机忙", "请稍后重试"),
        ],
        [1000, 3500, 4860],
    )

    add_heading(doc, "7. 三个必须掌握的流程", 1)
    add_heading(doc, "7.1 主机缓存仍新鲜：控制室直接取温度", 2)
    add_code_block(doc, [
        "G -> M1：READ_TEMP，request_id=100，允许缓存",
        "M1 -> G：TEMP_36，request_id=100，来源=主机缓存",
    ])
    add_heading(doc, "7.2 主机缓存过期：主机向同组从机取温度", 2)
    add_code_block(doc, [
        "G  -> M1：READ_TEMP，request_id=100，允许缓存",
        "M1 -> S1：READ_TEMP，request_id=100，强制新采样",
        "S1 -> M1：TEMP_36，request_id=100，来源=从机新采样",
        "M1 -> G ：TEMP_36，request_id=100，来源=从机新采样",
    ])
    add_heading(doc, "7.3 设置 M2 频率为 50.00 Hz", 2)
    add_code_block(doc, [
        "G  -> M2：CMD，SET_FREQ，value=5000",
        "M2 -> G ：ACK，已收到",
        "M2 -> TD710：写频率",
        "TD710 -> M2：正确 Modbus 回包",
        "M2 -> G ：RESULT，成功，当前频率=5000",
    ])
    add_callout(doc, "绝对禁止", "只收到 ACK 就在上位机显示“设置成功”。只有 RESULT 的结果码为 00 才能显示成功。", color=RED, fill="FDEEEE")

    add_heading(doc, "8. 重发、去重与异常规则", 1)
    add_table(
        doc,
        ["场景", "统一处理方式"],
        [
            ("TEMP_36 未收到", "请求方超时后重新发送同一个 READ_TEMP；不拆分补发单个温度点。"),
            ("CMD 未收到 ACK", "控制室以相同 seq 重发 CMD；重发次数必须有限。"),
            ("重复 CMD", "主机不能重复驱动 TD710；直接回复此前保存的 ACK 或 RESULT。"),
            ("CRC 错误", "直接丢弃，不回复，等待请求方超时重发。"),
            ("组号或目标不匹配", "直接丢弃，不得执行。"),
            ("从机超时", "主机向控制室发送 ERROR 或 TEMP_36 不完整状态；不得伪造正常温度。"),
            ("温度不完整或无效", "上位机明确显示无效；自动停机条件不成立。"),
        ],
        [2800, 6560],
    )
    add_body(doc, "超时时间和最大重发次数不是写死在本协议中的常量，必须在联调时依据 LoRa 实际速率、距离和丢包情况配置；但所有重发必须有限次、带超时、不得阻塞其他调度。")

    add_heading(doc, "9. 模块配置与上线检查", 1)
    add_table(
        doc,
        ["检查项", "统一要求", "验证现象"],
        [
            ("无线信道", "全部模块相同", "G、M、S 能互相接收指定报文"),
            ("LoRa 包长", "全部模块 128 B", "94 B 的 TEMP_36 单包完整收发"),
            ("串口参数", "沿用工程确认的 115200、8N1、无流控", "模块正常进入透传并收发二进制数据"),
            ("组号/拨码", "M1-S1、M2-S2、M3-S3、M4-S4 分别一致", "不同组请求不会得到错误组回复"),
            ("单包温度", "36 点一次 TEMP_36", "上位机一次更新 36 张温度卡片"),
            ("命令结果", "ACK + RESULT", "TD710 无回包时上位机显示失败而非成功"),
        ],
        [2000, 4050, 3310],
    )

    add_heading(doc, "10. 实施验收清单", 1)
    add_body(doc, "□ 从机未被主机点名时，不向 LoRa 主动发温度。")
    add_body(doc, "□ 控制室请求 M1 时，只有 M1 可以回复；M2-M4 不得回复。")
    add_body(doc, "□ 一条 TEMP_36 能让上位机一次更新当前主机的 36 个温度卡片。")
    add_body(doc, "□ 温度为 0 的点显示为“--”，且不参与自动停机。")
    add_body(doc, "□ 手动停止后，即使温度超限，风机仍保持停止；收到 SET_AUTO 后才恢复自动判断。")
    add_body(doc, "□ 频率、目标温度、模式掉电后仍保留；Flash 或 TD710 失败必须返回失败 RESULT。")
    add_body(doc, "□ 4 组同信道运行时，M4 或 S4 离线不影响其余三组轮询。")

    add_heading(doc, "11. 文档状态与边界", 1)
    add_body(doc, "本文件定义 LoRa 应用层报文，不替代 WH-L101-L 的 AT 参数说明，也不替代 TD710 的 Modbus 寄存器说明。模块实际地址、LoRa 速率、发送超时和自动重传次数应在现场联调参数表中统一记录。")
    add_body(doc, "旧的“从机每秒按 Port 上报 ASCII 文本”的遥测格式自本协议生效日起废弃。新开发、修改和联调必须只使用本协议。")
    add_body(doc, "依据：项目已确认的最终控制方案；WH-L101-L 配套点对点协议支持 32 B、64 B、128 B、240 B LoRa 包长。本系统统一采用 128 B。")


def main():
    doc = Document()
    set_styles(doc)
    add_title_page(doc)
    add_main_content(doc)
    props = doc.core_properties
    props.title = "智能通风系统 LoRa 报文协议（最终版）"
    props.subject = "控制室网关、主机、从机统一 LoRa 应用层报文格式"
    props.author = "智能通风系统项目组"
    props.comments = "V1.0 最终协议"
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
